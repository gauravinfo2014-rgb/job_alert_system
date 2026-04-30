"""
Generates a tailored Word (.docx) resume for a specific job.
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_section_heading(doc, text):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
    # horizontal line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2563EB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def generate_resume(job: dict, analysis: dict, output_dir: str = ".") -> str:
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    # Header
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GAURAV")
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"UX Designer / Product Designer  |  {job['company']} Application")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("gaurav.info2014@gmail.com  •  Portfolio: Available on request")
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph()

    # ATS score banner
    score = analysis.get("ats_score", 0)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    cells = t.rows[0].cells
    set_cell_bg(cells[0], "EFF6FF")
    cells[0].text = f"ATS Match: {score}/100"
    cells[0].paragraphs[0].runs[0].bold = True
    cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(cells[1], "EFF6FF")
    strengths = analysis.get("strengths", [])
    cells[1].text = "✓ " + "  ✓ ".join(strengths[:2]) if strengths else ""
    cells[1].paragraphs[0].runs[0].font.size = Pt(9)
    set_cell_bg(cells[2], "EFF6FF")
    gaps = analysis.get("gaps", [])
    cells[2].text = "⚠ " + "  ⚠ ".join(gaps[:2]) if gaps else "No significant gaps"
    cells[2].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Summary
    add_section_heading(doc, "Professional Summary")
    summary = analysis.get("tailored_summary",
        f"Experienced UX/Product Designer applying for {job['title']} at {job['company']}.")
    sp = doc.add_paragraph(summary)
    for run in sp.runs: run.font.size = Pt(10)

    # Skills
    add_section_heading(doc, "Core Skills")
    skills_data = [
        ["UX/UI Design", "User Research", "Prototyping", "Design Systems"],
        ["Figma", "Adobe XD", "Sketch", "InVision"],
        ["Wireframing", "Usability Testing", "A/B Testing", "Accessibility (WCAG)"],
        ["Interaction Design", "Visual Design", "Information Architecture", "Agile/Scrum"],
    ]
    st = doc.add_table(rows=len(skills_data), cols=4)
    for i, row_data in enumerate(skills_data):
        for j, skill in enumerate(row_data):
            cell = st.rows[i].cells[j]
            cell.text = f"• {skill}"
            for para in cell.paragraphs:
                for run in para.runs: run.font.size = Pt(9)

    # Experience
    add_section_heading(doc, "Professional Experience")
    experiences = [
        {"title": "Senior UX Designer", "company": "Recent Company", "period": "2022 – Present",
         "bullets": ["Led end-to-end design for key product features, increasing user engagement by 35%",
                     "Established and maintained a comprehensive design system used across 4 product teams",
                     "Conducted 50+ user research sessions translating insights into actionable improvements",
                     "Collaborated with engineering in Agile environment to ship features on schedule"]},
        {"title": "UX Designer", "company": "Previous Company", "period": "2019 – 2022",
         "bullets": ["Designed mobile-first experiences for iOS and Android apps with 1M+ users",
                     "Reduced onboarding drop-off rate by 28% through iterative UX improvements",
                     "Created high-fidelity prototypes and ran A/B tests to validate design hypotheses",
                     "Presented design concepts to C-suite stakeholders"]},
        {"title": "UI/UX Designer", "company": "Earlier Company", "period": "2017 – 2019",
         "bullets": ["Designed responsive web interfaces for e-commerce and SaaS products",
                     "Built wireframes and interactive prototypes using Figma and Adobe XD",
                     "Worked closely with developers to ensure pixel-perfect implementation"]},
    ]
    for exp in experiences:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(6)
        tr = tp.add_run(exp["title"])
        tr.bold = True; tr.font.size = Pt(10)
        tr.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
        cp = doc.add_paragraph()
        cp.paragraph_format.space_after = Pt(2)
        cr = cp.add_run(f"{exp['company']}  |  {exp['period']}")
        cr.font.size = Pt(9); cr.italic = True
        cr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
        for bullet in exp["bullets"]:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Inches(0.2)
            bp.paragraph_format.space_after = Pt(1)
            bp.add_run(bullet).font.size = Pt(9)

    # Education
    add_section_heading(doc, "Education")
    ep = doc.add_paragraph()
    er = ep.add_run("Bachelor of Design / Bachelor of Technology")
    er.bold = True; er.font.size = Pt(10)
    ep2 = doc.add_paragraph("University Name  |  Graduation Year")
    for run in ep2.runs:
        run.font.size = Pt(9); run.italic = True
        run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    # Tailoring tips footnote
    doc.add_paragraph()
    tips_para = doc.add_paragraph()
    tips_run = tips_para.add_run(
        f"── AI Tips for {job['company']} ──\n" +
        "\n".join(f"• {t}" for t in analysis.get("tailoring_tips", [])))
    tips_run.font.size = Pt(8); tips_run.italic = True
    tips_run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

    safe_co = job["company"].replace(" ", "_").replace("/", "-")
    safe_ti = job["title"].replace(" ", "_").replace("/", "-")[:30]
    filepath = os.path.join(output_dir, f"Resume_Gaurav_{safe_co}_{safe_ti}.docx")
    doc.save(filepath)
    print(f"  Resume saved: {filepath}")
    return filepath
